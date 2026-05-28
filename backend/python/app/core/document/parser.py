# -*- coding: utf-8 -*-
"""
炼丹炉 - 文档解析模块 (Document Parser)
解读各类丹方文件，提取其中真言，犹如解读上古丹经

支持的丹方格式:
    - Word (.doc, .docx) - python-docx
    - Excel (.xls, .xlsx) - openpyxl
    - Markdown (.md) - 直接读取
    - Text (.txt) - 直接读取
    - PDF (.pdf) - pdfplumber
"""
import os
import re
import logging
from abc import ABC, abstractmethod
from typing import Dict, Any, List, Optional

# 设置炼丹日志
logger = logging.getLogger(__name__)


# ==================== 解析器基类 ====================

class BaseParser(ABC):
    """
    丹方解析器基类 - 所有解析器之祖
    
    犹如丹方解读之基础法门，各类丹方解析器皆从此衍生。
    """
    
    def __init__(self, file_path: str) -> None:
        """
        初始化解析器
        
        Args:
            file_path: 丹方文件路径
        """
        self.file_path = file_path
        self.file_name = os.path.basename(file_path)
    
    @abstractmethod
    def parse(self) -> str:
        """
        解析丹方，提取文本
        
        Returns:
            提取的完整文本
        
        Raises:
            FileNotFoundError: 丹方文件不存在
            PermissionError: 无权读取丹方
            Exception: 解析失败
        """
        pass
    
    @abstractmethod
    def parse_with_metadata(self) -> List[Dict[str, Any]]:
        """
        解析丹方，提取带元数据的文本块
        
        Returns:
            文本块列表，每块包含 content 和 metadata
        """
        pass
    
    def _check_file_exists(self) -> None:
        """检查丹方文件是否存在 - 验丹方之真伪"""
        if not os.path.exists(self.file_path):
            logger.error(f"丹方文件不存在: {self.file_path}")
            raise FileNotFoundError(f"丹方文件不存在: {self.file_path}")
        if not os.path.isfile(self.file_path):
            logger.error(f"路径非文件: {self.file_path}")
            raise ValueError(f"路径非文件: {self.file_path}")


# ==================== Word 文档解析器 ====================

class DocxParser(BaseParser):
    """
    Word 文档解析器 - 解读字诀丹方
    
    使用 python-docx 解析 .docx 文件，提取段落文本。
    对于 .doc 文件，提示用户转换为 .docx 格式。
    """
    
    def parse(self) -> str:
        """解析 Word 文档，提取全部文本"""
        self._check_file_exists()
        
        # 若为 .doc 格式，提示转换
        if self.file_path.lower().endswith(".doc"):
            raise ValueError(
                "「.doc」格式为上古丹方，请转换为「.docx」后再行解读"
            )
        
        try:
            import docx
            logger.info(f"开始解读字诀丹方: {self.file_name}")
            
            document = docx.Document(self.file_path)
            paragraphs: List[str] = []
            
            for para in document.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            result = "\n".join(paragraphs)
            logger.info(f"字诀丹方解读完毕: {self.file_name}, "
                       f"共 {len(paragraphs)} 段, {len(result)} 字")
            return result
            
        except ImportError:
            logger.error("python-docx 未安装，无法解读字诀丹方")
            raise ImportError("请安装 python-docx: pip install python-docx")
        except Exception as e:
            logger.error(f"字诀丹方解读失败: {self.file_name}, 错误: {e}")
            raise RuntimeError(f"字诀丹方解读失败: {e}")
    
    def parse_with_metadata(self) -> List[Dict[str, Any]]:
        """解析 Word 文档，按段落提取带元数据的文本块"""
        self._check_file_exists()
        
        if self.file_path.lower().endswith(".doc"):
            raise ValueError("「.doc」格式为上古丹方，请转换为「.docx」后再行解读")
        
        try:
            import docx
            logger.info(f"开始分段解读字诀丹方: {self.file_name}")
            
            document = docx.Document(self.file_path)
            chunks: List[Dict[str, Any]] = []
            
            for idx, para in enumerate(document.paragraphs):
                text = para.text.strip()
                if text:
                    chunks.append({
                        "content": text,
                        "metadata": {
                            "file_name": self.file_name,
                            "file_type": "docx",
                            "paragraph_index": idx,
                            "style": para.style.name if para.style else "",
                        }
                    })
            
            logger.info(f"字诀丹方分段解读完毕: {len(chunks)} 段")
            return chunks
            
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")
        except Exception as e:
            logger.error(f"字诀丹方分段解读失败: {e}")
            raise RuntimeError(f"字诀丹方分段解读失败: {e}")


# ==================== Excel 解析器 ====================

class ExcelParser(BaseParser):
    """
    Excel 表格解析器 - 解读数诀丹方
    
    使用 openpyxl 解析 .xlsx 文件，每 sheet 为一页，
    逐行读取 cell 内容并拼接。
    """
    
    def parse(self) -> str:
        """解析 Excel 文件，提取全部文本"""
        self._check_file_exists()
        
        if self.file_path.lower().endswith(".xls"):
            raise ValueError("「.xls」格式为上古丹方，请转换为「.xlsx」后再行解读")
        
        try:
            import openpyxl
            logger.info(f"开始解读数诀丹方: {self.file_name}")
            
            workbook = openpyxl.load_workbook(self.file_path, data_only=True)
            all_texts: List[str] = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                all_texts.append(f"【{sheet_name}】")
                
                for row in sheet.iter_rows(values_only=True):
                    # 过滤空值，拼接行内 cell 文本
                    cell_texts = [str(cell).strip() for cell in row 
                                  if cell is not None and str(cell).strip()]
                    if cell_texts:
                        all_texts.append("\t".join(cell_texts))
            
            result = "\n".join(all_texts)
            logger.info(f"数诀丹方解读完毕: {self.file_name}, "
                       f"{len(workbook.sheetnames)} 页, {len(result)} 字")
            return result
            
        except ImportError:
            raise ImportError("请安装 openpyxl: pip install openpyxl")
        except Exception as e:
            logger.error(f"数诀丹方解读失败: {self.file_name}, 错误: {e}")
            raise RuntimeError(f"数诀丹方解读失败: {e}")
    
    def parse_with_metadata(self) -> List[Dict[str, Any]]:
        """解析 Excel 文件，按 sheet 提取带元数据的文本块"""
        self._check_file_exists()
        
        if self.file_path.lower().endswith(".xls"):
            raise ValueError("「.xls」格式为上古丹方，请转换为「.xlsx」后再行解读")
        
        try:
            import openpyxl
            logger.info(f"开始分页解读数诀丹方: {self.file_name}")
            
            workbook = openpyxl.load_workbook(self.file_path, data_only=True)
            chunks: List[Dict[str, Any]] = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                rows: List[str] = []
                
                for row in sheet.iter_rows(values_only=True):
                    cell_texts = [str(cell).strip() for cell in row 
                                  if cell is not None and str(cell).strip()]
                    if cell_texts:
                        rows.append("\t".join(cell_texts))
                
                if rows:
                    chunks.append({
                        "content": f"【{sheet_name}】\n" + "\n".join(rows),
                        "metadata": {
                            "file_name": self.file_name,
                            "file_type": "xlsx",
                            "sheet_name": sheet_name,
                            "row_count": len(rows),
                        }
                    })
            
            logger.info(f"数诀丹方分页解读完毕: {len(chunks)} 页")
            return chunks
            
        except ImportError:
            raise ImportError("请安装 openpyxl: pip install openpyxl")
        except Exception as e:
            logger.error(f"数诀丹方分页解读失败: {e}")
            raise RuntimeError(f"数诀丹方分页解读失败: {e}")


# ==================== Markdown 解析器 ====================

class MarkdownParser(BaseParser):
    """
    Markdown 文档解析器 - 解读符诀丹方
    
    直接读取 .md 文件内容。
    """
    
    def parse(self) -> str:
        """解析 Markdown 文件"""
        self._check_file_exists()
        
        try:
            logger.info(f"开始解读符诀丹方: {self.file_name}")
            
            with open(self.file_path, "r", encoding="utf-8") as f:
                content = f.read()
            
            logger.info(f"符诀丹方解读完毕: {self.file_name}, {len(content)} 字")
            return content
            
        except UnicodeDecodeError:
            # 尝试其他编码
            with open(self.file_path, "r", encoding="gbk") as f:
                content = f.read()
            return content
        except Exception as e:
            logger.error(f"符诀丹方解读失败: {self.file_name}, 错误: {e}")
            raise RuntimeError(f"符诀丹方解读失败: {e}")
    
    def parse_with_metadata(self) -> List[Dict[str, Any]]:
        """解析 Markdown 文件，按标题分段"""
        self._check_file_exists()
        
        try:
            logger.info(f"开始分段解读符诀丹方: {self.file_name}")
            
            with open(self.file_path, "r", encoding="utf-8") as f:
                content = f.read()
            
            # 按二级标题分段
            sections = re.split(r'\n(?=##\s)', content)
            chunks: List[Dict[str, Any]] = []
            
            for idx, section in enumerate(sections):
                section = section.strip()
                if section:
                    # 提取标题
                    title_match = re.match(r'^#+\s*(.+)$', section, re.MULTILINE)
                    title = title_match.group(1) if title_match else f"第{idx+1}段"
                    
                    chunks.append({
                        "content": section,
                        "metadata": {
                            "file_name": self.file_name,
                            "file_type": "md",
                            "section_index": idx,
                            "section_title": title,
                        }
                    })
            
            logger.info(f"符诀丹方分段解读完毕: {len(chunks)} 段")
            return chunks
            
        except Exception as e:
            logger.error(f"符诀丹方分段解读失败: {e}")
            raise RuntimeError(f"符诀丹方分段解读失败: {e}")


# ==================== 文本解析器 ====================

class TextParser(BaseParser):
    """
    纯文本解析器 - 解读简诀丹方
    
    直接读取 .txt 文件内容。
    """
    
    def parse(self) -> str:
        """解析文本文件"""
        self._check_file_exists()
        
        try:
            logger.info(f"开始解读简诀丹方: {self.file_name}")
            
            with open(self.file_path, "r", encoding="utf-8") as f:
                content = f.read()
            
            logger.info(f"简诀丹方解读完毕: {self.file_name}, {len(content)} 字")
            return content
            
        except UnicodeDecodeError:
            with open(self.file_path, "r", encoding="gbk") as f:
                content = f.read()
            return content
        except Exception as e:
            logger.error(f"简诀丹方解读失败: {self.file_name}, 错误: {e}")
            raise RuntimeError(f"简诀丹方解读失败: {e}")
    
    def parse_with_metadata(self) -> List[Dict[str, Any]]:
        """解析文本文件，按行提取"""
        self._check_file_exists()
        
        try:
            logger.info(f"开始分段解读简诀丹方: {self.file_name}")
            
            with open(self.file_path, "r", encoding="utf-8") as f:
                lines = f.readlines()
            
            chunks: List[Dict[str, Any]] = []
            for idx, line in enumerate(lines):
                text = line.strip()
                if text:
                    chunks.append({
                        "content": text,
                        "metadata": {
                            "file_name": self.file_name,
                            "file_type": "txt",
                            "line_number": idx + 1,
                        }
                    })
            
            logger.info(f"简诀丹方分段解读完毕: {len(chunks)} 行")
            return chunks
            
        except Exception as e:
            logger.error(f"简诀丹方分段解读失败: {e}")
            raise RuntimeError(f"简诀丹方分段解读失败: {e}")


# ==================== PDF 解析器 ====================

class PDFParser(BaseParser):
    """
    PDF 文档解析器 - 解读图诀丹方
    
    使用 pdfplumber 解析 PDF 文件，逐页提取文本。
    """
    
    def parse(self) -> str:
        """解析 PDF 文件，提取全部文本"""
        self._check_file_exists()
        
        try:
            import pdfplumber
            logger.info(f"开始解读图诀丹方: {self.file_name}")
            
            all_texts: List[str] = []
            
            with pdfplumber.open(self.file_path) as pdf:
                logger.info(f"图诀丹方共 {len(pdf.pages)} 页")
                
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text and text.strip():
                        all_texts.append(f"--- 第 {page_num} 页 ---\n{text.strip()}")
            
            result = "\n\n".join(all_texts)
            logger.info(f"图诀丹方解读完毕: {self.file_name}, "
                       f"{len(pdf.pages)} 页, {len(result)} 字")
            return result
            
        except ImportError:
            logger.error("pdfplumber 未安装，无法解读图诀丹方")
            raise ImportError("请安装 pdfplumber: pip install pdfplumber")
        except Exception as e:
            logger.error(f"图诀丹方解读失败: {self.file_name}, 错误: {e}")
            raise RuntimeError(f"图诀丹方解读失败: {e}")
    
    def parse_with_metadata(self) -> List[Dict[str, Any]]:
        """解析 PDF 文件，按页提取带元数据的文本块"""
        self._check_file_exists()
        
        try:
            import pdfplumber
            logger.info(f"开始分页解读图诀丹方: {self.file_name}")
            
            chunks: List[Dict[str, Any]] = []
            
            with pdfplumber.open(self.file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text and text.strip():
                        chunks.append({
                            "content": text.strip(),
                            "metadata": {
                                "file_name": self.file_name,
                                "file_type": "pdf",
                                "page_number": page_num,
                                "total_pages": len(pdf.pages),
                            }
                        })
            
            logger.info(f"图诀丹方分页解读完毕: {len(chunks)} 页")
            return chunks
            
        except ImportError:
            raise ImportError("请安装 pdfplumber: pip install pdfplumber")
        except Exception as e:
            logger.error(f"图诀丹方分页解读失败: {e}")
            raise RuntimeError(f"图诀丹方分页解读失败: {e}")


# ==================== 解析器工厂 ====================

class ParserFactory:
    """
    丹方解析器工厂 - 万法归一
    
    根据丹方类型（文件扩展名）自动选择对应的解析器，
    犹如炼丹之万能丹炉，各类丹材皆可炼化。
    
    Usage:
        parser = ParserFactory.get_parser("/path/to/file.docx", "docx")
        text = parser.parse()
    """
    
    # 丹方类型到解析器的映射 - 丹方目录
    _parsers: Dict[str, type] = {
        "docx": DocxParser,
        "doc": DocxParser,    # .doc 会提示转换
        "xlsx": ExcelParser,
        "xls": ExcelParser,   # .xls 会提示转换
        "md": MarkdownParser,
        "txt": TextParser,
        "pdf": PDFParser,
    }
    
    @classmethod
    def get_parser(cls, file_path: str, file_type: Optional[str] = None) -> BaseParser:
        """
        获取对应类型的解析器 - 选取合适之丹炉
        
        Args:
            file_path: 丹方文件路径
            file_type: 文件类型（可选，不传则从路径推断）
        
        Returns:
            对应的解析器实例
        
        Raises:
            ValueError: 不支持的文件类型
        """
        if file_type is None:
            # 从文件路径推断类型
            file_type = os.path.splitext(file_path)[1].lstrip(".").lower()
        else:
            file_type = file_type.lower().lstrip(".")
        
        logger.info(f"丹方类型识别: {file_type}, 路径: {file_path}")
        
        if file_type not in cls._parsers:
            supported = ", ".join(cls._parsers.keys())
            logger.error(f"不支持的丹方类型: {file_type}")
            raise ValueError(
                f"此丹方类型「{file_type}」无法解读，"
                f"仅支持: {supported}"
            )
        
        parser_class = cls._parsers[file_type]
        logger.info(f"选定解析器: {parser_class.__name__}")
        return parser_class(file_path)
    
    @classmethod
    def supported_types(cls) -> List[str]:
        """获取支持的文件类型列表 - 查看可炼之丹材"""
        return list(cls._parsers.keys())
    
    @classmethod
    def register_parser(cls, file_type: str, parser_class: type) -> None:
        """
        注册新的解析器 - 新增炼丹之法
        
        Args:
            file_type: 文件类型
            parser_class: 解析器类（须继承 BaseParser）
        """
        if not issubclass(parser_class, BaseParser):
            raise ValueError("解析器必须继承 BaseParser")
        cls._parsers[file_type.lower()] = parser_class
        logger.info(f"新解析器注册成功: {file_type} -> {parser_class.__name__}")
